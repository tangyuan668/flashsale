import sys
import shutil
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

TEMPLATE_PATH = r'D:\桌面\毕设\202211040450-史佳楠-基于Redis的高并发实时秒杀系统的设计与实现-本科毕业设计（论文）中期报告1111.docx'

# Backup
backup_path = TEMPLATE_PATH + '.bak'
shutil.copy2(TEMPLATE_PATH, backup_path)
print(f"备份已保存: {backup_path}")

doc = Document(TEMPLATE_PATH)

# ============================================================
# 1. Save template references BEFORE any modifications
# ============================================================
s2_heading = doc.paragraphs[9]    # "2.1 微服务架构搭建" - bold, Times New Roman
s2_content = doc.paragraphs[14]   # Content paragraph - not bold
s2_empty = doc.paragraphs[12]     # Empty paragraph

s3_subheading = doc.paragraphs[46]  # "3.1 防重复下单与幂等控制问题" - bold
s3_content = doc.paragraphs[47]     # "问题描述：..." - not bold

# Save section 3 XML elements (paragraphs 44-66, total 23 elements)
s3_elements = [doc.paragraphs[i]._element for i in range(44, 67)]

# Clone diagram (65) and caption (66) from original section 3
diagram_clone = deepcopy(s3_elements[21])  # paragraph 65 - Consolas diagram
caption_clone = deepcopy(s3_elements[22])  # paragraph 66 - "图3 异步补偿机制架构图"

s3_header = s3_elements[0]  # Section 3 header element

# ============================================================
# 2. Helper functions
# ============================================================

def make_new_paragraph(template_para, text):
    """Clone template paragraph's formatting and set new text."""
    new_p = deepcopy(template_para._element)
    # Remove all existing runs
    for r in new_p.findall(qn('w:r')):
        new_p.remove(r)
    if text:
        new_r = OxmlElement('w:r')
        # Copy run properties from template's first run
        if template_para.runs:
            rPr = template_para.runs[0]._element.find(qn('w:rPr'))
            if rPr is not None:
                new_r.append(deepcopy(rPr))
        new_t = OxmlElement('w:t')
        new_t.text = text
        new_t.set(qn('xml:space'), 'preserve')
        new_r.append(new_t)
        new_p.append(new_r)
    return new_p


def set_element_text(element, template_para, text):
    """Modify existing paragraph element's text, keeping paragraph properties."""
    for r in element.findall(qn('w:r')):
        element.remove(r)
    if text:
        new_r = OxmlElement('w:r')
        if template_para.runs:
            rPr = template_para.runs[0]._element.find(qn('w:rPr'))
            if rPr is not None:
                new_r.append(deepcopy(rPr))
        new_t = OxmlElement('w:t')
        new_t.text = text
        new_t.set(qn('xml:space'), 'preserve')
        new_r.append(new_t)
        element.append(new_r)


# ============================================================
# 3. Build new section 2 content (2.8, 2.9, 2.10)
# ============================================================

s2_additions = []

# --- 2.8 高并发场景下的幂等控制机制 ---
s2_additions.append(make_new_paragraph(s2_empty, ""))
s2_additions.append(make_new_paragraph(s2_heading, "2.8 高并发场景下的幂等控制机制"))
s2_additions.append(make_new_paragraph(s2_content,
    "在高并发秒杀场景中，幂等控制是保障系统正确性的关键。系统在请求层和消息消费层两个维度设计了幂等校验机制，有效防止了因网络抖动、页面重复提交或消息重复投递导致的重复操作。"
))
s2_additions.append(make_new_paragraph(s2_content,
    '在库存扣减层面，系统在Redis的Lua扣减脚本中设计了双层幂等校验。第一层通过"deducting:{userId}:{stockKey}"键设置扣减中标记，有效期设为120秒，在窗口期内同一用户对同一商品的重复请求会被直接拦截，返回"正在处理中"的提示，防止网络抖动或页面重复提交导致的重复扣减；第二层通过"user:purchase:{userId}:{stockKey}"键记录用户已购买标记，在订单创建成功后设置，用于永久防止同一用户重复购买。两层标记协同工作，短时间防重提交与长效购买幂等控制互补，幂等控制的粒度为用户+商品维度，不同商品之间互不影响，用户可以同时下单多个不同商品。'
))
s2_additions.append(make_new_paragraph(s2_content,
    '在消息消费层面，系统在消费者端引入了Redis去重键机制。以库存消费者为例，每次消费消息前，先检查"d:inventory:{orderNo}"（操作成功标记）和"d:inventory:processing:{orderNo}"（处理中标记，30分钟有效期）两个键。若已存在成功标记则直接返回，表示该消息已被处理；若存在处理中标记则说明消息正在被另一个节点处理，同样跳过。只有两个键均不存在时才执行实际操作，并在操作完成后设置成功标记。对于可重试的异常（如网络超时、连接拒绝等），系统会删除处理中标记允许立即重试；对于不可重试的异常，保留处理中标记等待30分钟自动过期。该机制与RocketMQ的最大重试间隔（30分钟）相匹配，确保在消息重试窗口期内不会出现重复处理。'
))

# --- 2.9 数据一致性保障与补偿机制 ---
s2_additions.append(make_new_paragraph(s2_empty, ""))
s2_additions.append(make_new_paragraph(s2_heading, "2.9 数据一致性保障与补偿机制"))
s2_additions.append(make_new_paragraph(s2_content,
    "在异步消息驱动的微服务架构中，数据一致性是核心挑战。系统通过本地消息表、定时补偿任务和回滚机制，构建了一套覆盖全链路的一致性保障体系。"
))
s2_additions.append(make_new_paragraph(s2_content,
    '在消息可靠性方面，系统在订单服务和库存服务中均引入了本地消息表机制。在发送MQ消息前，先将消息内容通过独立事务（Propagation.REQUIRES_NEW）持久化到本地数据库的消息表中，消息记录写入成功后再尝试发送MQ消息。若发送成功则更新消息状态为"已发送"；若发送失败，消息记录保留为"待发送"状态，由定时任务每隔30秒扫描一次待发送消息，采用指数退避策略（30秒、1分钟、2分钟、4分钟、8分钟）进行最多5次重试。消息认领过程使用CAS机制防止多节点重复认领：定时任务将消息状态从"待发送"（status=0）更新为"发送中"（status=3），只有成功抢占的节点才执行实际发送，确保每条消息仅被一个节点处理。'
))
s2_additions.append(make_new_paragraph(s2_content,
    '在订单超时处理方面，系统设置了15分钟的支付超时阈值，通过定时任务每分钟扫描超时未支付的订单。取消操作采用CAS乐观锁（基于订单status字段），将status=0（待支付）更新为status=2（已取消），仅当状态符合预期时才执行更新。订单取消成功后，通过本地消息表发送库存回滚MQ消息，触发库存服务通过专用的ROLLBACK_STOCK_LUA脚本原子性地完成库存恢复、扣减标记清除和购买记录清除，同时清理Redis中的用户购买标记，使用户可以重新参与秒杀。'
))
s2_additions.append(make_new_paragraph(s2_content,
    '在库存补偿方面，系统设计了库存补偿定时任务，每分钟扫描库存变动日志中超过10分钟仍未关联到有效订单的库存记录。对于每条待补偿记录，系统通过Feign远程调用查询订单服务，判断原始订单是否已存在：若订单已存在则跳过补偿；若订单不存在，则进一步查询该用户对同一商品是否有后续成功重试的订单。若用户已通过重试成功下单，则仅回滚库存但保留购买标记（通过COMPENSATE_ROLLBACK_STOCK_ONLY_LUA脚本），防止误删导致用户可以重复购买；若用户无后续成功订单，则执行完整回滚（通过COMPENSATE_ROLLBACK_LUA脚本，恢复库存+清除扣减标记+清除购买标记），用户可重新发起秒杀。补偿完成后设置"order:compensated"标记（24小时有效期），防止延迟到达的库存结果消息仍然创建订单。所有Feign调用在失败时均采用保守处理策略——默认认为订单存在或有成功订单，宁可跳过补偿也不误回滚。'
))
s2_additions.append(make_new_paragraph(s2_content,
    '综上所述，系统通过Redis幂等标记、本地消息表、CAS乐观锁、Lua原子脚本以及多层时间窗口的补偿任务，构建了一套覆盖"请求→处理→消息→确认→补偿"全链路的防重与兜底机制，在保证高并发性能的同时有效维护了分布式环境下的数据一致性。异步补偿机制的整体架构如下图所示：'
))

# Diagram and caption (cloned from original section 3)
s2_additions.append(diagram_clone)
s2_additions.append(caption_clone)

# --- 2.10 分布式环境下的并发控制策略 ---
s2_additions.append(make_new_paragraph(s2_empty, ""))
s2_additions.append(make_new_paragraph(s2_heading, "2.10 分布式环境下的并发控制策略"))
s2_additions.append(make_new_paragraph(s2_content,
    "在系统多实例部署场景下，多个节点可能同时处理同一订单的取消、同一消息的发送或同一库存记录的补偿，若不加以控制可能导致重复操作或数据不一致。系统在数据库层面广泛采用CAS（Compare-And-Swap）乐观锁策略，所有状态变更操作均以当前状态作为WHERE条件，仅当状态符合预期时才执行更新，返回影响行数为0则表示已被其他节点处理，当前节点直接跳过。该策略应用于订单取消（status=0到2）、消息认领（status=0到3, retry_count匹配）、库存补偿（基于操作时间戳判断）等多个关键流程，有效避免了分布式环境下的并发冲突。相比于重量级的分布式锁，CAS乐观锁具有更低的性能开销，更适合高并发场景。"
))

# Insert all new section 2 paragraphs before section 3 header (in order)
for elem in s2_additions:
    s3_header.addprevious(elem)

print(f"已在第二部分末尾插入 {len(s2_additions)} 个段落（2.8/2.9/2.10）")

# ============================================================
# 4. Replace section 3 content with new problems & solutions
# ============================================================

# New section 3 content (17 items -> s3_elements[1] through s3_elements[17])
# Element 0 (header) stays unchanged: "三、存在的问题和拟解决方法"
new_s3_items = [
    (s3_content,
     "在系统核心功能开发完成后的自测与联调过程中，发现以下几类需要进一步解决的问题："),
    (s2_empty, ""),
    (s3_subheading, "3.1 支付功能尚未集成"),
    (s3_content,
     '问题描述：目前系统实现了从用户发起秒杀请求到订单创建的完整流程，但订单创建后缺少实际的支付环节。订单状态停留在"待支付"（status=0），用户无法通过系统完成支付操作，导致订单只能通过15分钟超时机制自动取消，完整的秒杀业务闭环尚未形成。'),
    (s3_content,
     '拟解决方法：计划在下一阶段集成支付宝沙箱环境或实现模拟支付功能，完成订单支付接口的开发。考虑到毕业设计的演示需求，拟优先实现模拟支付方案：用户在订单详情页点击"模拟支付"按钮后，系统将订单状态更新为"已支付"（status=1）并记录支付时间，以此验证支付完成后订单状态流转和库存确认等后续业务逻辑的正确性。'),
    (s2_empty, ""),
    (s3_subheading, "3.2 库存数据持久化与一致性校验待完善"),
    (s3_content,
     '问题描述：当前系统为追求极致的库存扣减性能，将库存数据完全存储在Redis中，秒杀过程中的库存变动未实时同步回数据库。虽然Redis自身的持久化机制（RDB/AOF）可以在一定程度上防止数据丢失，但在Redis发生严重故障的场景下，数据库中的库存数据可能与实际不一致。此外，系统重启后缺少自动化的库存数据校验与恢复机制。'),
    (s3_content,
     '拟解决方法：计划在Redis库存扣减成功后，通过MQ消息异步更新数据库中的库存记录，实现Redis与数据库的最终一致性。同时，在系统启动时增加库存数据校验逻辑，对比Redis缓存与数据库中的库存数据，在数据不一致时以数据库为准重新预热，确保系统重启后的数据正确性。'),
    (s2_empty, ""),
    (s3_subheading, "3.3 系统测试覆盖不足"),
    (s3_content,
     '问题描述：目前系统尚未编写单元测试和集成测试，核心业务逻辑（如库存扣减Lua脚本、库存补偿判断逻辑、消费者幂等控制等）的正确性主要依赖手动测试验证，难以全面覆盖各种边界场景和异常情况，系统的可靠性缺乏自动化测试保障。'),
    (s3_content,
     '拟解决方法：计划使用JUnit 5和Mockito框架为核心业务模块编写单元测试，重点覆盖库存扣减逻辑、消息消费幂等逻辑、库存补偿判断逻辑等关键路径。同时，使用JMeter进行并发压力测试，模拟高并发秒杀场景，验证系统在并发环境下的正确性和稳定性，根据测试结果进行针对性的优化调整。'),
    (s2_empty, ""),
    (s3_subheading, "3.4 配置管理与安全性优化"),
    (s3_content,
     '问题描述：当前系统中JWT密钥、数据库密码等敏感信息以硬编码方式写在源代码和配置文件中，存在安全隐患，不符合安全编码规范。此外，用户密码的加密存储目前为简化实现，未使用标准的BCrypt算法进行哈希处理。'),
    (s3_content,
     '拟解决方法：计划将JWT密钥、数据库密码等敏感配置迁移至Nacos配置中心进行统一管理，通过环境变量或加密配置的方式避免敏感信息出现在代码仓库中。同时，引入Spring Security提供的BCryptPasswordEncoder替换当前的简化加密实现，确保用户密码以标准的BCrypt哈希格式存储，提升系统的安全性。'),
]

# Modify section 3 elements (skip element 0 = header, keep unchanged)
for i, (template, text) in enumerate(new_s3_items):
    element_index = i + 1
    set_element_text(s3_elements[element_index], template, text)

# Remove excess section 3 elements (18-22, total 5 elements to remove)
removed_count = 0
for elem in s3_elements[18:]:
    elem.getparent().remove(elem)
    removed_count += 1

print(f"已替换第三部分内容，移除了 {removed_count} 个多余段落")

# ============================================================
# 5. Save
# ============================================================
doc.save(TEMPLATE_PATH)
print(f"\n文件已保存: {TEMPLATE_PATH}")
print("修改完成！")
