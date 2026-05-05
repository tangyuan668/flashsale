# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 设置页面边距
for section in doc.sections:
    section.top_margin = Cm(2.8)
    section.bottom_margin = Cm(2.8)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)


def add_heading_text(text, font_size=12, bold=True, space_after=6, space_before=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = Pt(22)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = 'Times New Roman'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p


def add_body_text(text, indent=True, space_after=3, space_before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = Pt(22)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p


def add_code_block(text, space_after=6, space_before=3):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = Pt(16)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.name = 'Consolas'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p


def add_figure_caption(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(22)
    p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = 'Times New Roman'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_table(headers, rows, caption=''):
    if caption:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = Pt(22)
        run = p.add_run(caption)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.name = 'Times New Roman'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10.5)
                run.font.name = 'Times New Roman'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = val
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10.5)
                    run.font.name = 'Times New Roman'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_paragraph()  # spacing after table


# ==================== 正文内容 ====================

add_body_text('.', indent=False)
add_body_text('.', indent=False)

add_heading_text('一、任务书中本阶段工作目标与任务要求')

add_body_text(
    '根据任务书的进度安排，本阶段（第4—8周）的核心工作目标为实现系统核心功能模块的开发，'
    '重点围绕高并发秒杀场景下的关键技术展开。具体任务要求包括：（1）完成系统微服务架构的搭建与各服务间的通信配置；'
    '（2）实现基于Redis的库存预热与Lua脚本原子扣减逻辑，确保并发场景下库存数据的正确性，杜绝超卖现象；'
    '（3）集成Redisson分布式锁，对库存预热、活动状态切换等关键流程进行互斥控制；'
    '（4）完成用户注册登录、秒杀活动管理等基础功能模块的开发；'
    '（5）实现订单服务的异步下单流程，引入RocketMQ消息队列进行流量削峰；'
    '（6）确保秒杀主链路（用户请求→库存扣减→订单创建→库存确认）的完整跑通。'
)

add_heading_text('二、目前已完成任务情况')

add_body_text(
    '截至目前，本课题已基本完成第4—8周的核心功能开发任务，系统主链路已初步跑通。'
    '系统的整体架构如下图所示，采用基于Spring Cloud Alibaba的微服务架构，将系统拆分为用户服务、活动服务、'
    '库存服务、订单服务和API网关五个核心模块，各服务通过Nacos进行注册与配置管理，通过Feign进行服务间远程调用，'
    '通过RocketMQ进行异步消息通信。'
)

add_code_block(
    '                         用户客户端\n'
    '                              |\n'
    '                              v\n'
    '                         Nginx 反向代理\n'
    '                              |\n'
    '                              v\n'
    '  +----------------------------------------------------------+\n'
    '  |                   API Gateway (8080)                     |\n'
    '  |              路由转发 / JWT认证 / 限流 / 防刷              |\n'
    '  +------+----------+----------+----------+------------------+\n'
    '         |          |          |          |\n'
    '         v          v          v          v\n'
    '    +---------+ +---------+ +---------+ +---------+\n'
    '    |  User   | |Activity | |Inventory| |  Order  |\n'
    '    | (8081)  | | (8082)  | | (8083)  | | (8084)  |\n'
    '    +---------+ +---------+ +---------+ +---------+\n'
    '         |          |          |          |\n'
    '         v          v          v          v\n'
    '      用户库     活动库     库存库     订单库\n'
    '         |          |          |          |\n'
    '         +----------+----------+----------+\n'
    '                    |          |\n'
    '          +---------+    +----------+\n'
    '          |  Nacos  |    |  Redis   |\n'
    '          +---------+    +----------+\n'
    '                    |\n'
    '              +-----------+\n'
    '              | RocketMQ  |\n'
    '              +-----------+',
    space_before=6
)

add_figure_caption('图1 系统总体架构图')

add_body_text('具体完成情况如下：')

# ---- 2.1 ----
add_heading_text('2.1 微服务架构搭建', bold=True, space_before=6)

add_body_text(
    '基于Spring Boot 3.2与Spring Cloud Alibaba完成了系统整体架构的搭建，'
    '将系统拆分为用户服务、活动服务、库存服务、订单服务和API网关五个核心模块。'
    '各服务均已注册至Nacos配置中心，实现了服务间的Feign远程调用与统一配置管理。'
    '同时，各服务采用独立数据库的部署方式，通过分库策略实现数据的物理隔离，降低服务间的耦合度。'
)

add_table(
    ['模块名称', '端口', '职责说明', '独立数据库'],
    [
        ['flashsale-gateway', '8080', 'API网关，路由转发、JWT认证、限流、防刷', '--'],
        ['flashsale-user', '8081', '用户注册登录、JWT令牌管理', 'flashsale_user'],
        ['flashsale-activity', '8082', '秒杀活动管理、库存预热', 'flashsale_activity'],
        ['flashsale-inventory', '8083', '库存扣减与回滚、库存补偿', 'flashsale_inventory'],
        ['flashsale-order', '8084', '异步下单、订单管理、超时取消', 'flashsale_order'],
        ['flashsale-common', '--', '公共模块，统一响应、异常处理、工具类', '--'],
    ],
    '表1 系统模块划分'
)

# ---- 2.2 ----
add_heading_text('2.2 用户服务与认证', bold=True, space_before=6)

add_body_text(
    '完成了用户注册、登录功能，采用BCrypt算法对用户密码进行加密存储，有效防止密码明文泄露风险。'
    '登录成功后系统生成JWT令牌返回给客户端，实现无状态认证机制。网关层集成了JWT鉴权过滤器（AuthenticationFilter），'
    '对非公开接口进行统一的身份校验，将用户ID和手机号注入请求头传递给下游服务，下游服务无需重复解析令牌。'
)

add_body_text(
    '此外，网关层还集成了令牌桶限流过滤器（RateLimitFilter）和防刷过滤器（AntiBrushFilter），'
    '共同构成系统的入口防护体系。令牌桶限流基于Redis+Lua脚本实现，对每个用户的请求速率进行精确控制；'
    '防刷过滤器对同一用户的请求频率进行实时监控，超过阈值自动将用户IP加入黑名单。'
)

# ---- 2.3 ----
add_heading_text('2.3 库存预热与原子扣减', bold=True, space_before=6)

add_body_text(
    '在活动服务中实现了库存预热功能。管理员在秒杀活动开始前，通过接口触发库存预热操作。'
    '预热过程中，系统使用Redisson分布式锁对同一活动的预热操作进行互斥控制，防止多节点部署情况下的重复预热。'
    '预热逻辑将活动商品的库存数据从数据库加载至Redis缓存中，以stock:cache:{activityId}:{itemId}为键存储库存数量，'
    '为后续的高并发库存扣减提供数据基础。'
)

add_body_text(
    '库存扣减采用Redis+Lua脚本实现，这是本系统防超卖的核心机制。在单次Lua脚本执行中原子性地完成以下四项操作：'
    '第一，检查Redis中该商品的库存是否充足，库存不足直接返回失败；'
    '第二，检查用户是否已购买过该商品（通过user:purchase键），已购买则拒绝；'
    '第三，检查是否存在正在进行的扣减操作（通过deducting键，有效期120秒），防止窗口期内的重复请求；'
    '第四，执行库存扣减并设置扣减中标记。Lua脚本的核心逻辑如下：'
)

add_code_block(
    '-- DEDUCT_STOCK_LUA 核心逻辑\n'
    'local stock = tonumber(redis.call("GET", key))\n'
    'if stock < quantity then return -1 end            -- 库存不足\n'
    'if redis.call("EXISTS", purchaseKey) == 1\n'
    '    then return -4 end                             -- 已购买\n'
    'if redis.call("EXISTS", deductingKey) == 1\n'
    '    then return -3 end                             -- 扣减中\n'
    'redis.call("SET", key, stock - quantity)          -- 扣减库存\n'
    'redis.call("SETEX", deductingKey, 120, orderNo)   -- 设置扣减中标记\n'
    'return stock - quantity',
    space_before=3
)

add_body_text(
    '该Lua脚本由Redis单线程执行，保证了在高并发场景下多个请求对同一商品的库存操作是串行化的，'
    '从根本上杜绝了超卖问题。扣减中标记（deducting键）的有效期设为120秒，在订单创建成功后由订单服务清除；'
    '若120秒内未收到确认，则由库存补偿定时任务处理。'
)

# ---- 2.4 ----
add_heading_text('2.4 分布式锁控制', bold=True, space_before=6)

add_body_text(
    '集成Redisson客户端，在库存预热和活动状态切换等关键流程中使用分布式锁进行互斥控制。'
    '以库存预热为例，系统通过Redisson获取以活动ID为维度的分布式锁，设置10秒等待超时和30秒锁持有超时，'
    '确保同一活动的预热操作在同一时刻只被一个服务节点执行，防止多节点部署情况下的重复预热或数据不一致。'
    '锁的获取与释放均在try-finally块中完成，确保异常情况下锁能够正确释放。'
)

# ---- 2.5 ----
add_heading_text('2.5 异步下单与消息队列', bold=True, space_before=6)

add_body_text(
    '引入RocketMQ实现下单流程的异步化处理。系统采用"异步双MQ"模式：'
    '用户发起秒杀请求后，订单服务生成订单号并将请求信息保存到本地消息表，'
    '随后通过MQ消息通知库存服务执行库存扣减；库存服务完成扣减后，'
    '再通过MQ消息将结果通知回订单服务，由订单服务完成实际的订单创建。'
    '这种设计将同步的下单操作拆分为多个异步步骤，用户请求到达后立即返回"处理中"状态，'
    '前端通过轮询获取最终结果，有效避免了同步等待带来的连接占用问题。系统的秒杀核心流程如下图所示：'
)

add_code_block(
    '用户发起秒杀\n'
    '    |\n'
    '    v\n'
    'Gateway (JWT认证 / 令牌桶限流 / IP防刷)\n'
    '    |\n'
    '    v\n'
    '订单服务: 前置幂等检查 (purchase键 / deducting键)\n'
    '    |\n'
    '    v\n'
    '生成订单号 + 保存本地消息表\n'
    '    |\n'
    '    v\n'
    '发送MQ (SECKILL_ORDER_TOPIC) --> 库存服务\n'
    '    |\n'
    '    v\n'
    '库存消费者 (消息去重: d:inventory键)\n'
    '    |\n'
    '    v\n'
    'Redis + Lua 原子扣减库存\n'
    '    |-- 失败 --> 发送失败MQ --> 订单服务 --> 记录失败状态\n'
    '    |-- 成功 --> 写库存日志 --> 发送结果MQ (STOCK_RESULT_TOPIC)\n'
    '                                    |\n'
    '                                    v\n'
    '                    订单消费者 (消息去重: d:order键)\n'
    '                                    |\n'
    '                    检查补偿标记 / 购买标记\n'
    '                        |-- 已补偿 --> 拒绝创建\n'
    '                        |-- 已购买 --> 发送回滚MQ\n'
    '                        |-- 通过   --> 创建订单 + 设置Redis标记',
    space_before=3
)

add_figure_caption('图2 秒杀核心流程图')

# ---- 2.6 ----
add_heading_text('2.6 本地消息表与消息可靠性', bold=True, space_before=6)

add_body_text(
    '为解决消息发送失败导致的数据一致性问题，在订单服务和库存服务中均引入了本地消息表机制。'
    '在发送MQ消息前，先将消息内容通过独立事务（Propagation.REQUIRES_NEW）持久化到本地数据库的消息表中，'
    '消息记录写入成功后再尝试发送MQ消息。若发送成功则更新消息状态为"已发送"；'
    '若发送失败，消息记录保留为"待发送"状态，由定时任务每隔30秒扫描一次待发送消息，'
    '采用指数退避策略（30秒、1分钟、2分钟、4分钟、8分钟）进行最多5次重试。'
)

add_body_text(
    '消息认领过程使用CAS机制防止多节点重复认领：定时任务将消息状态从"待发送"（status=0）'
    '更新为"发送中"（status=3），只有成功抢占的节点才执行实际发送，确保每条消息仅被一个节点处理。'
    '本地消息表的状态流转如下表所示：'
)

add_table(
    ['status值', '状态含义', '说明'],
    [
        ['0', '待发送', '初始状态，或发送失败后恢复'],
        ['3', '发送中', 'CAS抢占后的临时状态，防止多节点重复发送'],
        ['1', '已发送', 'MQ消息发送成功'],
        ['2', '发送失败', '超过最大重试次数'],
    ],
    '表2 本地消息表状态流转'
)

# ---- 2.7 ----
add_heading_text('2.7 数据库设计', bold=True, space_before=6)

add_body_text(
    '系统采用分库策略，每个微服务拥有独立的数据库，共包含5个数据库：用户库（flashsale_user）、'
    '活动库（flashsale_activity）、库存库（flashsale_inventory）、订单库（flashsale_order）。'
    '此外，订单服务和库存服务各部署一份本地消息表（local_message），用于保证消息的可靠投递。'
    '各服务核心数据表结构如下：'
)

add_table(
    ['字段名', '类型', '约束', '备注'],
    [
        ['id', 'bigint', '主键，自增', '用户ID'],
        ['username', 'varchar(50)', '唯一', '用户名'],
        ['phone', 'varchar(20)', '唯一', '手机号'],
        ['password', 'varchar(255)', '非空', '密码（BCrypt加密）'],
        ['status', 'tinyint', '默认1', '状态（0禁用，1正常）'],
        ['create_time', 'datetime', '--', '创建时间'],
    ],
    '表3 用户表（user）'
)

add_table(
    ['字段名', '类型', '约束', '备注'],
    [
        ['id', 'bigint', '主键，自增', '活动ID'],
        ['name', 'varchar(100)', '非空', '活动名称'],
        ['start_time', 'datetime', '非空', '开始时间'],
        ['end_time', 'datetime', '非空', '结束时间'],
        ['status', 'tinyint', '默认0', '0待开始 1进行中 2已结束 3已取消'],
        ['preheat_status', 'tinyint', '默认0', '预热状态（0未预热，1已预热）'],
    ],
    '表4 活动表（activity）'
)

add_table(
    ['字段名', '类型', '约束', '备注'],
    [
        ['id', 'bigint', '主键，自增', '库存ID'],
        ['activity_id / item_id', 'bigint', '外键', '活动ID与商品ID'],
        ['total_stock', 'int', '非空', '总库存'],
        ['available_stock', 'int', '非空', '可用库存'],
        ['version', 'int', '默认0', '乐观锁版本号'],
    ],
    '表5 库存表（inventory）'
)

add_table(
    ['字段名', '类型', '约束', '备注'],
    [
        ['id', 'bigint', '主键，自增', '订单ID'],
        ['order_no', 'varchar(64)', '唯一', '订单号（雪花算法）'],
        ['user_id', 'bigint', '外键', '用户ID'],
        ['total_amount', 'decimal(10,2)', '非空', '订单总金额'],
        ['status', 'tinyint', '默认0', '0待支付 1已支付 2已取消 3已超时'],
        ['cancel_reason', 'varchar(255)', '--', '取消原因'],
    ],
    '表6 订单主表（order_info）'
)

add_table(
    ['字段名', '类型', '约束', '备注'],
    [
        ['id', 'bigint', '主键，自增', '消息ID'],
        ['business_no', 'varchar(64)', '--', '业务编号（订单号）'],
        ['topic', 'varchar(100)', '非空', 'MQ主题'],
        ['message_body', 'text', '非空', '消息体（JSON）'],
        ['status', 'tinyint', '默认0', '0待发送 1已发送 2失败 3发送中'],
        ['retry_count', 'int', '默认0', '已重试次数'],
        ['max_retry', 'int', '默认5', '最大重试次数'],
        ['next_retry_time', 'datetime', '--', '下次重试时间'],
    ],
    '表7 本地消息表（local_message）'
)

# ==================== 第三部分 ====================

add_heading_text('三、存在的问题和拟解决方法')

add_body_text(
    '在核心功能的开发与联调过程中，围绕分布式环境下数据一致性与系统可靠性，'
    '发现并重点解决了以下几类问题：'
)

# ---- 3.1 ----
add_heading_text('3.1 防重复下单与幂等控制问题', bold=True, space_before=6)

add_body_text('问题描述：在高并发场景下，用户可能因网络抖动、页面重复提交等原因发起多次秒杀请求，若不加以控制将导致重复扣减库存。')

add_body_text(
    '解决方案：系统在Redis的Lua扣减脚本中设计了多层幂等校验机制。'
    '首先通过"deducting:{userId}:{stockKey}"键设置扣减中标记，有效期设为120秒，'
    '在120秒窗口期内同一用户对同一商品的重复请求会被直接拦截，返回"正在处理中"的提示；'
    '其次通过"user:purchase:{userId}:{stockKey}"键记录用户已购买标记，'
    '在订单创建成功后设置，用于永久防止同一用户重复购买。'
    '两层标记协同工作，既保证了短时间内的防重提交，又实现了长效的购买幂等控制。'
    '其中{stockKey}为"stock:cache:{activityId}:{itemId}"，是Redis中存储该商品库存数量的键，'
    '幂等控制的粒度为用户+商品维度，不同商品之间互不影响，用户可以同时下单多个不同商品。'
)

# ---- 3.2 ----
add_heading_text('3.2 消息发送失败与可靠性问题', bold=True, space_before=6)

add_body_text(
    '问题描述：在异步下单流程中，用户发起秒杀请求后，订单服务生成订单号并向库存服务发送MQ消息，'
    '由库存服务执行Redis库存扣减，扣减完成后通过MQ通知订单服务创建订单。'
    '在这一过程中，若发送MQ消息时消息队列不可用或网络异常，消息将丢失，导致整个下单流程无法继续。'
)

add_body_text(
    '解决方案：系统引入了本地消息表机制。在发送MQ消息前，先将消息内容持久化到本地消息表中，'
    '保存操作使用独立事务（Propagation.REQUIRES_NEW）确保消息记录可靠写入；'
    '消息记录写入成功后再尝试发送MQ消息。若发送成功则更新消息状态为"已发送"；'
    '若发送失败，消息记录保留为"待发送"状态，由定时任务每隔30秒扫描一次待发送消息，'
    '采用指数退避策略（30秒、1分钟、2分钟、4分钟、8分钟）进行最多5次重试。'
    '消息认领过程使用CAS机制（基于status和retry_count字段以及"发送中"临时状态）防止多节点重复认领，'
    '确保每条消息仅被一个节点处理。此外，标记消息为"已发送"的操作同样使用独立事务，'
    '确保状态更新不受外层异常影响，避免消息已发出但因状态回滚而被定时任务重复发送。'
)

# ---- 3.3 ----
add_heading_text('3.3 消息消费幂等与重复消息处理问题', bold=True, space_before=6)

add_body_text('问题描述：RocketMQ在某些异常情况下可能对同一消息进行重复投递，消费者端若不做幂等处理将导致库存被重复操作。')

add_body_text(
    '解决方案：系统在消费者端引入了Redis去重键机制。以库存消费者为例，每次消费消息前，'
    '先通过Lua脚本原子性地检查"d:inventory:{orderNo}"（库存操作成功标记）和'
    '"d:inventory:processing:{orderNo}"（库存处理中标记，30分钟有效期）两个键。'
    '若已存在成功标记则直接返回，表示该消息已被处理；若存在处理中标记则说明消息正在被另一个节点处理，同样跳过。'
    '只有两个键均不存在时才执行实际的库存操作，并在操作完成后设置成功标记。'
    '对于可重试的异常（如网络超时、连接拒绝等），系统会删除处理中标记允许立即重试；'
    '对于不可重试的异常，保留处理中标记等待30分钟自动过期。'
    '该机制与RocketMQ的最大重试间隔（30分钟）相匹配，确保在消息重试窗口期内不会出现重复处理。'
)

# ---- 3.4 ----
add_heading_text('3.4 订单超时未支付与库存回滚问题', bold=True, space_before=6)

add_body_text('问题描述：用户成功下单后若未在规定时间内完成支付，系统需自动取消订单并回滚库存。若取消操作不当，可能导致库存无法正确恢复。')

add_body_text(
    '解决方案：系统设置了15分钟的支付超时阈值，通过定时任务每分钟扫描一次超时未支付的订单。'
    '取消操作采用CAS乐观锁（基于订单status字段），将status=0（待支付）更新为status=2（已取消），'
    '仅当状态符合预期时才执行更新，返回影响行数为0则表示已被其他节点处理。'
    '订单取消成功后，通过本地消息表发送库存回滚MQ消息，触发库存服务执行回滚操作。'
    '回滚操作通过专用的ROLLBACK_STOCK_LUA脚本原子性地完成库存恢复、扣减标记清除和购买记录清除。'
    '同时清理Redis中的用户购买标记（user:purchase键），使用户可以重新参与秒杀。'
    '本地消息表保证了回滚消息的可靠投递，即使发送失败也会由定时任务自动重试。'
)

# ---- 3.5 ----
add_heading_text('3.5 库存扣减后订单创建失败的一致性问题', bold=True, space_before=6)

add_body_text(
    '问题描述：在秒杀流程中，Redis库存已扣减但订单创建可能因数据库异常等原因失败，'
    '导致库存被"悬挂"占用，用户既无法获得订单也无法重新购买。'
)

add_body_text(
    '解决方案：系统设计了库存补偿机制。库存补偿定时任务每分钟执行一次，'
    '扫描库存变动日志中超过10分钟仍未关联到有效订单的库存记录。'
    '对于每条待补偿记录，系统首先通过Feign远程调用查询订单服务，判断原始订单是否已存在：'
    '若订单已存在则跳过补偿；若订单不存在，则进一步查询该用户对同一商品是否有后续成功重试的订单。'
    '若用户已通过重试成功下单，则仅回滚库存但保留购买标记（通过COMPENSATE_ROLLBACK_STOCK_ONLY_LUA脚本），'
    '防止误删导致用户可以重复购买；若用户无后续成功订单，则执行完整回滚'
    '（通过COMPENSATE_ROLLBACK_LUA脚本，恢复库存+清除扣减标记+清除购买标记），用户可重新发起秒杀。'
    '补偿完成后设置"order:compensated"标记（24小时有效期），防止延迟到达的库存结果消息仍然创建订单。'
    '所有Feign调用在失败时均采用保守处理策略——默认认为订单存在或有成功订单，宁可跳过补偿也不误回滚。'
)

# ---- 3.6 ----
add_heading_text('3.6 分布式环境下多节点并发操作的冲突问题', bold=True, space_before=6)

add_body_text(
    '问题描述：在系统多实例部署场景下，多个节点可能同时处理同一订单的取消、'
    '同一消息的发送或同一库存记录的补偿，若不加以控制可能导致重复操作或数据不一致。'
)

add_body_text(
    '解决方案：系统在数据库层面广泛采用CAS（Compare-And-Swap）乐观锁策略。'
    '所有状态变更操作均以当前状态作为WHERE条件，仅当状态符合预期时才执行更新，'
    '返回影响行数为0则表示已被其他节点处理，当前节点直接跳过。'
    '该策略应用于订单取消（status=0到2）、消息认领（status=0到3, retry_count匹配）、'
    '库存补偿（基于操作时间戳判断）等多个关键流程，有效避免了分布式环境下的并发冲突。'
    '相比于重量级的分布式锁，CAS乐观锁具有更低的性能开销，更适合高并发场景。'
)

# 总结段
add_body_text(
    '综上所述，系统通过Redis幂等标记、本地消息表、CAS乐观锁、Lua原子脚本以及多层时间窗口的补偿任务，'
    '构建了一套覆盖"请求→处理→消息→确认→补偿"全链路的防重与兜底机制，'
    '在保证高并发性能的同时有效维护了分布式环境下的数据一致性。'
    '异步补偿机制的整体架构如下图所示：'
)

add_code_block(
    '消息可靠性:                    库存补偿:                    订单超时取消:\n'
    '发送失败                        订单创建失败                   超时未支付\n'
    '  |                               |                            |\n'
    '  v                               v                            v\n'
    '本地消息表                     补偿定时任务                   超时扫描任务\n'
    'status=0                       每分钟执行                     每分钟执行\n'
    '  |                               |                            |\n'
    '  v                               v                            v\n'
    'CAS抢占重试                    Feign查订单                   CAS取消订单\n'
    '30s->1m->2m->4m->8m             |--存在->跳过                 status=0->2\n'
    '最多5次                        |--有重试->仅回滚库存               |\n'
    '                               |--无重试->完整回滚                v\n'
    '                                                          本地消息表\n'
    '                                                          发回滚MQ',
    space_before=3
)

add_figure_caption('图3 异步补偿机制架构图')

# ==================== 第四部分 ====================

add_heading_text('四、下一步工作计划')

add_body_text(
    '在本阶段核心功能开发完成的基础上，下一阶段（第9-13周）的主要工作计划如下：'
)

add_body_text(
    '（1）系统集成与联调优化。对各服务模块进行全面的集成测试，修复联调过程中发现的逻辑缺陷，'
    '确保秒杀主链路在各种场景下的稳定运行。'
)

add_body_text(
    '（2）并发压力测试与性能优化。使用JMeter等工具构造一定规模的并发请求模拟秒杀场景，'
    '对系统的响应时间、吞吐量以及资源使用情况进行观察和分析。根据测试结果对系统进行优化调整，'
    '例如优化Redis连接池配置、调整限流参数、优化数据库查询等。'
)

add_body_text(
    '（3）前端页面开发。开发秒杀活动列表页、商品详情页、秒杀下单页、订单查询页等前端页面，'
    '与后端接口进行对接，完成用户交互流程的闭环。'
)

add_body_text(
    '（4）毕业论文撰写。对系统设计与实现过程进行整理归纳，开始撰写毕业论文，'
    '系统阐述设计思路、关键技术实现与测试结果。'
)

# ==================== 签字部分 ====================

add_body_text('', indent=False, space_before=12)
add_body_text('指导教师签字：', indent=False)

add_body_text('', indent=False, space_before=24)
add_body_text('    年  月  日', indent=False)

add_body_text('', indent=False, space_before=12)
add_heading_text('五、系级教学单位审核意见：')

add_body_text('', indent=False, space_before=12)
add_body_text('中期考核分数：          ', indent=False)

add_body_text('', indent=False, space_before=12)
add_body_text('考核组长签字：', indent=False)

add_body_text('', indent=False, space_before=24)
add_body_text('    年  月  日', indent=False)

# ==================== 保存 ====================

output_path = r'D:\桌面\毕设\202211040450-史佳楠-基于Redis的高并发实时秒杀系统的设计与实现-本科毕业设计（论文）中期报告.docx'
doc.save(output_path)
print(f'文件已保存成功: {output_path}')
